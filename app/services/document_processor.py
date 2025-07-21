import os 
import PyPDF2
from docx import Document
import pytesseract
from PIL import Image
from typing import List, Dict, Any
import fitz

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.process_pdf,
            '.docx': self.process_docx,
            '.doc': self.process_doc,
            '.txt': self.process_txt,
            '.jpg': self.process_image,
            '.png': self.process_image,
            '.jpeg': self.process_image,
        }
    
    def process_document(self, file_path: str) -> dict:
        """Process the uploaded document and extract text from it"""
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        processor = self.supported_formats[file_extension]
        content = processor(file_path)

        return {
            'file_name': os.path.basename(file_path),
            'file_extension': file_extension,
            'content': content,
            'file_path': file_path,
            'file_type': file_extension,
            'metadata': self.metadata(file_path)
        }
    
    def process_pdf(self, file_path: str) -> str:
        """Process PDF files and extract text"""
        text = ""
        try:
            # First try using fitz (PyMuPDF)
            docs = fitz.open(file_path)
            for page in docs:
                text += page.get_text()
            docs.close()
        except:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text()
        return text
    
    def process_docx(self, file_path: str) -> str:
        """Process DOCX files and extract text"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def process_doc(self, file_path: str) -> str:
        """Process DOC files and extract text"""
        # For .doc files, we'll try to convert them or use a fallback
        try:
            # Try to open as docx first (if it's actually a docx file with wrong extension)
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except:
            # Fallback: return a message that .doc files need conversion
            return f"Note: .doc files need to be converted to .docx format for proper processing. File: {os.path.basename(file_path)}"
    
    def process_txt(self, file_path: str) -> str:
        """Process TXT files and extract text"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    
    def process_image(self, file_path: str) -> str:
        """Process image files and extract text using OCR"""
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text
    
    def metadata(self, file_path: str) -> dict:
        """Extract metadata from the file"""
        return {
            'file_name': os.path.basename(file_path),
            'file_extension': os.path.splitext(file_path)[1].lower(),
            'file_size': os.path.getsize(file_path),
            'file_path': file_path,
            'file_type': os.path.splitext(file_path)[1].lower(),
        }


